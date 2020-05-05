
# ---------------------- REQUIREMENTS ---------------------- #


# Documentation: https://python-docx.readthedocs.io/en/latest/
import docx
from docx import Document
import xlsxwriter
import datetime
import os



# ---------------------- READ FILE ---------------------- #



# Get input file from user
# clippings_file = input("\nPlease enter path to clippings file: \n------------------------------------------\n")
clippings_file = "My Clippings.txt"

# Open file
text = open(clippings_file, mode="r", encoding ='utf-8-sig')

# Create list containing each individual line
lines = text.readlines()

# Close file
text.close()



# ---------------------- USEFUL FUNCTIONS ---------------------- #



def get_title(section):
    title = section[0].split("(", 1)[0].strip()
    return title

def get_authors(section):
    authors = section[0].split("(", 1)[1].replace(")","").strip().split(";")
    return authors

def get_location(section):
    location = section[1].split("Added on", 1)[0].replace("- Your Highlight ","").strip()
    if "at location" in location: 
        index = location.find("at location") + 12
    elif "on page" in location:
        index = location.find("location") + 9

    location = location[index:index + 10]
    if "-" in location: 
        location = location.split("-")[0]
    location = location.replace("|","").strip()
    
    return int(location)

def get_date(section):
    date = section[1].split("Added on", 1)[1].split(", ", 1)[1]
    return date

def get_quote(section):
    quote = section[3].strip()
    return quote



# ---------------------- 1. EXTRACT INFORMATION ---------------------- #



# List of dicts to store structured information 
sections = []

# Temp value to store current section
section = []

# Iterate over each line
for line in lines:

    # If separator is found, extract information from section and reset section
    if "==========" in line:
        information = {
            "title": get_title(section),
            "authors" : get_authors(section),
            "location" : get_location(section),
            "date" : get_date(section),
            "quote" : get_quote(section),
        }      

        # Add structured information to list of dicts
        sections.append(information)

        # Reset temp section
        section = []
        
            
    # Else continue adding lines to the existing section
    else:
        section.append(line)

# Order list of quotes based on title and location
sections = sorted(sections, key = lambda i: (i["title"] , i['location'])) 



### -------------------------------------- 2. WRITE ALL QUOTES TO EXCEL -------------------------------------- ###

today = datetime.date.today()

try:
    os.mkdir("clips_" + str(today))
except:
    print("<!> Folder could not be created.\n")
else: 
    print("<!> Folder was successfully created.\n")
workbook = xlsxwriter.Workbook('clips_' + str(today) + '/#clippings_ ' + str(today) + '.xlsx') 
worksheet = workbook.add_worksheet('Clippings') 

# Set filter on headers
worksheet.autofilter('A1:E1')
  
# Set the columns widths
worksheet.set_column('A:A', 30)
worksheet.set_column('B:B', 20)
worksheet.set_column('C:C', 15)
worksheet.set_column('D:D', 120)
worksheet.set_column('E:E', 30)

# Set row height
worksheet.set_row(0, 20)

# Add a header format 
header = workbook.add_format()
header.set_bold(1)
header.set_font_color('white')
header.set_bg_color('#4E4E4E')
header.set_align('center')
header.set_align('vcenter')

# Add a cell format
cell_format = workbook.add_format()
cell_format.set_text_wrap()
cell_format.set_align('center')
cell_format.set_align('vcenter')
 
# Add an Excel date format.
date_format = workbook.add_format({'num_format': 'mmmm d yyyy'})
date_format.set_text_wrap()
date_format.set_align('center')
date_format.set_align('vcenter')

# Add a wrap format.
wrap = workbook.add_format({'num_format': 'mmmm d yyyy'})
wrap.set_text_wrap()
wrap.set_align('center')
wrap.set_align('left')
 
# Write some data headers.
worksheet.write('A1', 'Book Title', header)
worksheet.write('B1', 'Location', header)
worksheet.write('C1', 'Date', header)
worksheet.write('D1', 'Quote', header)
worksheet.write('E1', 'Authors', header)

# Start from the first cell. Rows and columns are zero indexed. 
row = 1
col = 0

# Iterate over the data and write it out row by row. 
for section in sections: 
    worksheet.write(row, col, section["title"], cell_format) 
    worksheet.write(row, col + 1, section["location"], cell_format)

    # Convert the date string into a datetime object.
    worksheet.write(row, col + 2, section["date"], date_format)
    worksheet.write(row, col + 3, section["quote"], wrap)
    worksheet.write(row, col + 4, str(section["authors"])[1:-1].replace("'", ""), wrap)
    row += 1

print(f"___ EXCEL ___ \n We created #{row-1} rows in the Excel file. You can find 'clippings.xlsx' in the current folder.\n")
workbook.close() 

all_books = []

# Get unique title values
for section in sections:
    if section["title"] in all_books:
        continue
    else:
        text = section["title"]
        all_books.append(text)



def list_books(all_books):
    
    # ---------------------- 3. SHOW USER LIST OF BOOKS ---------------------- #

    # Present list of books to user
    print("\nAvailable Books:\n---------------------------")

    for book in all_books:
        print("->", book)
        


def choose_books(all_books):

    # ---------------------- 4. USER SELECTS BOOK ---------------------- #

    # Get input from user
    title = input("\nEnter desired book title: \n-------------------------\n")

    # Find book titles matching user input
    selected_books = [book for book in all_books if title in book]

    # If no match found, signal error
    if len(selected_books) == 0:
        print("\n<!> We didn't find any books.\n")
        decision = "n"

    # If at least one match found, confirm list with user
    else:
        print("\nWe found the following books: \n------------------------------")
        for i in selected_books: 
            print("->", i)

        decision = input("\nAre you sure you want to select these books?(y/n)\n--------------------------------------------\n")

        # If user says yes, go ahead and create the file
        if decision.lower() == "y" or decision.lower() == "yes":
            write_books(title, selected_books)
        else:
            return False

def write_books(title, selected_books):

# ---------------------- 5. WRITE QUOTES FROM BOOK IN WORD DOCUMENT ---------------------- #

    # Create file
    document = docx.Document()
    
    # Write book title
    document.add_heading('Selected books:', level = 1)

    # Write list of selected books
    document.add_paragraph(selected_books , style='List Bullet')

    # Filter list of quotes based on title
    sections_filter = [section for section in sections if section["title"] in selected_books]

    # Write authors (if one book selected only)

    # Write header
    document.add_heading('Authors:', level = 1)

    # Write list of authors
    for author in sections_filter[0]["authors"]:
        document.add_paragraph(author, style='List Bullet')

    # formatting...
    document.add_paragraph("")
    document.add_heading('Highlights', level = 1)
    document.add_paragraph("")
    
    # Write list of quotes associated with selected book(s)
    for section in sections_filter:
        document.add_paragraph(section["quote"] + " (" + str(section["location"]) + ")")
        
        # formatting...
        document.add_paragraph("______________________________")
        document.add_paragraph("")

    # Save file to disk
    document.save("clips_" + str(today) + "/"+title+".docx")

    # Confirm success to user
    print("\n<!> File was successfully created: '{0}.docx'".format(title))


    # Ask if he wishes to start over
    #choice = input("\nWould you like to try another book?(y/n)\n--------------------------------------------\n")
    #if choice.lower() == "y" or choice.lower() == "yes":
    #    choose_books(all_books)
    #else:
    #    exit

# Show all books to user
#list_books(all_books)

# Ask user to select book
#choose_books(all_books)


for book in all_books:
    title = book.replace(":", "").strip()[0:56]
    title = "".join(title)
    write_books(title, book)

